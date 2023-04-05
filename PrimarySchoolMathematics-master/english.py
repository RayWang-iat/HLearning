import random
import time
import json
import os
from docx import Document  # 引入docx类生成docx文档
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PrintPreview:
    '''本类负责生成完整的口算题文档使之适合打印机打印。可以生成多套题，生成数可以控。

    - @p_list   list
    需要打印口算题库，至少包含一套口算题

    - @p_title   list
    页面标题，这个标题的生成依据程序题型的选择和数字的范围选择而生成，例如：选择了0-20，加减法，进退位
    则自动生成标题为：0到20加减法进退位混合口算题，list中包含了多套题的页面标题名称

    - @p_column  int
    打印页排版口算题的列数

    '''

    def __init__(self, l, l2, l3, l4, m1, m2, m3, m4, tit, subtitle, col, tsize=26, subsize=11, csize=16,
                 # 默认输出文件地址为项目根目录
                 docxpath=os.path.join(os.path.dirname(os.path.abspath(__file__)), "")):
        '''
        :param l: list 需要打印的口算题列表
        :param tit: list 口算页标题
        :param subtitle str 小标题
        :param col: int 列数
        :param tsize: int 标题字号
        :param csize: int 口算题字号
        :param docxpath str 保存路径

        '''
        self.p_list = l
        self.p_list2 = l2
        self.p_list3 = l3
        self.p_list4 = l4
        self.p_m1 = m1
        self.p_m2 = m2
        self.p_m3 = m3
        self.p_m4 = m4
        self.p_title = tit
        self.p_subtitle = subtitle
        self.p_column = col
        self.p_title_size = tsize
        self.p_subtitle_size = subsize
        self.p_content_siae = csize
        self.docxpath = docxpath

    def create_psmdocx(self, l, title, docxname):
        '''
        :param l list 一组题库
        :param title str 页面标题
        :param docxname  str 题库保存文件名
        :return: none
        '''
        if (title == ''):
            page_title = '小学生口算题'
        else:
            page_title = title
        p_docx = Document()  # 创建一个docx文档
        p_docx.styles['Normal'].font.name = u'Times'  # 可换成word里面任意字体
        p = p_docx.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        run = p.add_run(page_title)
        run.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(self.p_title_size)  # 字体大小设置，和word里面的字号相对应

        sp = p_docx.add_paragraph()
        sp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        srun = sp.add_run(self.p_subtitle)
        srun.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun.font.size = Pt(self.p_subtitle_size)  # 字体大小设置，和word里面的字号相对应

        if self.p_m1 == 1:
            # 判断需要用到的行数
            if (len(l) % self.p_column):
                rs = len(l) // self.p_column + 2
            else:
                rs = len(l) // self.p_column + 1

            # 将口算题添加到docx表格中
            k = 0  # 计数器
            table = p_docx.add_table(rows=rs, cols=self.p_column)

            for i in range(rs):
                if i > 0:
                    row_cells = table.rows[i].cells
                    for j in range(self.p_column):
                        if (k > len(l) - 1):
                            break
                        else:
                            row_cells[j].text = l[k]
                            k = k + 1

        if self.p_m2 == 1:
            # 判断需要用到的行数
            rs = len(self.p_list2)+1

            # 将口算题添加到docx表格中
            k = 0  # 计数器
            table = p_docx.add_table(rows=rs, cols=1)

            for i in range(rs):
                if i > 0:
                    row_cells = table.rows[i].cells
                    for j in range(1):
                        if (k > len(self.p_list2) - 1):
                            break
                        else:
                            row_cells[j].text = self.p_list2[k]
                            k = k + 1

        if self.p_m3 == 1:
            # 判断需要用到的行数
            rs = len(self.p_list3)+1

            # 将口算题添加到docx表格中
            k = 0  # 计数器
            table = p_docx.add_table(rows=rs, cols=1)

            for i in range(rs):
                if i > 0:
                    row_cells = table.rows[i].cells
                    for j in range(1):
                        if (k > len(self.p_list3) - 1):
                            break
                        else:
                            row_cells[j].text = self.p_list3[k]
                            k = k + 1


        if self.p_m4 == 1:
            # 判断需要用到的行数
            rs = len(self.p_list4)+8

            # 将口算题添加到docx表格中
            k = 0  # 计数器
            table = p_docx.add_table(rows=rs, cols=1)

            for i in range(rs):
                if i > 0:
                    row_cells = table.rows[i].cells
                    for j in range(1):
                        if (k > len(self.p_list4) - 1):
                            break
                        else:
                            row_cells[j].text = self.p_list4[k]
                            k = k + 1

        table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        table.style.font.size = Pt(self.p_content_siae)  # 字体大小设置，和word里面的字号相对应
        print(self.docxpath + '{}.docx'.format(docxname))
        p_docx.save(self.docxpath + '{}.docx'.format(docxname))  # 输出docx

    def produce(self):
        k = 1
        for l, t in zip(self.p_list, self.p_title):
            self.create_psmdocx(l, t, t )
            k = k + 1

def no(str):
   return str!='None'

def get_random_list(start,stop,n):
    '''
    生成范围在[start,stop], 长度为n的数组.
    区间包含左右endpoint
    '''
    arr = list(range(start, stop+1))
    shuffle_n(arr,n)
    return arr[-n:]

def shuffle_n(arr,n):

    random.seed(time.time())
    for i in range(len(arr)-1,len(arr)-n-1,-1):
        j = random.randint(0,i)
        arr[i], arr[j] = arr[j], arr[i]


class EnglishGenerate(object):
    spelling=None  #单词拼写
    intertranslation=None #中英互译
    is_English=None #英语到中文（全是英语）
    is_Chinese=None #中文到英语（全是中文）
    is_Mixed=None #混合模式
    number=None

    def __init__(self,spelling, intertranslation, is_English, is_Chinese, is_Mixed, number,number2,number3,number4):
        if spelling is None and intertranslation is None:
            raise Exception("Error,Spelling or intertranslation shuold be choice")

        self.spelling=spelling
        self.intertranslation=intertranslation
        self.is_English=is_English
        self.is_Chinese=is_Chinese
        self.is_Mixed=is_Mixed
        self.number = number
        self.number2 = number2
        self.number3 = number3
        self.number4 = number4

    def get_data_from_txt(self):
        f = open("words.txt", encoding='utf-8')  # 返回一个文件对象
        line = f.readline()  # 调用文件的 readline()方法
        lst_e = []
        lst_c = []
        while line:
            txt = line.replace('\n', '')
            lst = txt.split('：')
            en = lst[0]
            ch = lst[1]
            lst_e.append(en)
            lst_c.append(ch)
            line = f.readline()
        f.close()

        randnum = random.randint(0, 100)
        random.seed(randnum)
        random.shuffle(lst_e)
        random.seed(randnum)
        random.shuffle(lst_c)

        lst_english=[]
        lst_chinese=[]
        i=0

        while i < self.number:
            lst_english.append(lst_e[i])
            lst_chinese.append(lst_c[i])
            i+=1

        return lst_english,lst_chinese

    def English_to_Chinese(self):
        (lst_english,lst_chinese)=self.get_data_from_txt()
        data_lst=['英译中','','']
        answer_lst=['英译中','','']
        i=0
        while i<self.number:
            s1 = "("+str(i+1)+")、"+lst_english[i] + ":__________"
            s2 = lst_english[i] + ":" + lst_chinese[i]
            data_lst.append(s1)
            answer_lst.append(s2)
            i += 1

        return data_lst

    def Chinese_to_English(self):
        (lst_english, lst_chinese) = self.get_data_from_txt()
        data_lst = ['中译英','','']
        answer_lst = ['中译英','','']
        i = 0
        while i<self.number:
            s1 = "("+str(i+1)+")、"+lst_chinese[i] + ":__________"
            s2 = lst_chinese[i] + ":" + lst_english[i]
            data_lst.append(s1)
            answer_lst.append(s2)
            i += 1

        return data_lst

    def Mixed(self):
        (lst_english, lst_chinese) = self.get_data_from_txt()
        i = 0
        data_lst = ['一、中英互译','','']
        answer_lst = ['一、中英互译','','']
        while i<self.number:
            flag = random.randint(1, 100)
            if flag%2==0:
                s1 = "("+str(i+1)+")、"+lst_chinese[i] + ":__________"
                s2 = lst_chinese[i] + ":" + lst_english[i]
                data_lst.append(s1)
                answer_lst.append(s2)
                i += 1
            else:
                s1 = "("+str(i+1)+")、"+lst_english[i] + ":__________"
                s2 = lst_english[i] + ":" + lst_chinese[i]
                data_lst.append(s1)
                answer_lst.append(s2)
                i += 1

        return data_lst

    def Spelling(self):
        f = open("words.txt", encoding='utf-8')  # 返回一个文件对象
        line = f.readline()  # 调用文件的 readline()方法
        lst_e = ['单词拼写','','']
        lst_c = ['单词拼写','','']
        while line:
            txt = line.replace('\n', '')
            lst = txt.split('：')
            en = lst[0]
            ch = lst[1]
            lst_e.append(en)
            lst_c.append(ch)
            line = f.readline()
        f.close()

        randnum = random.randint(0, 100)
        random.seed(randnum)
        random.shuffle(lst_e)
        random.seed(randnum)
        random.shuffle(lst_c)
        i=0
        q = 0
        data_lst=['单词拼写','','']
        answer_lst=[]

        while i < self.number:
            word = lst_e[q]
            chinese = lst_c[q]
            length = len(word)
            if length > 3:
                empty_lst = []
                first_random = random.randint(2, 6)
                if length > first_random:
                    presecond_random = get_random_list(1, length - 1, first_random)
                    second_random = sorted(presecond_random)
                    n = 0
                    m = 0
                    for j in word:
                        if m < first_random:
                            if n == second_random[m]:
                                empty_lst.append('_')
                                m += 1
                                n += 1
                            else:
                                empty_lst.append(j)
                                n += 1
                        else:
                            empty_lst.append(j)
                            n += 1
                    reword = "".join(empty_lst)
                    s1 = "("+str(i+1)+")、"+chinese + ":" + reword
                    s2 = "("+str(i+1)+")、"+chinese + ":" + word
                    data_lst.append(s1)
                    answer_lst.append(s2)
                    i += 1
                    q += 1
                else:
                    q += 1
            else:
                q += 1

        return data_lst

    def Choice_Question(self):
        number=self.number2
        while True:
            if number > 19:
                num1 = random.randint(0, 19)
                num2 = number - num1
            else:
                num1 = random.randint(0, number)
                num2 = number - num1
            if num1!=0 and num2!=0:
                break
        a = get_random_list(0, 18, num1)
        n1 = sorted(a)
        b = get_random_list(0, 99, num2)
        n2 = sorted(b)
        f = open("English_choice_1.json", encoding='utf-8')  # 返回一个文件对象
        line = json.load(f)  # 调用文件的 readline()方法
        length = len(line)
        i = 0
        c = 0
        k=0
        question_lst = ['选择填空部分']
        while i < length:
            if c < len(n1):
                if i == n1[c]:
                    data = line[i]
                    Question = data['Question']
                    Contents = [data['Content1'].replace("\xa0", ""), data['Content2'].replace("\xa0", " "),
                                data['Content3'].replace("\xa0", " "), data['Content4'].replace("\xa0", " "),
                                data['Content5'].replace("\xa0", " "), data['Content6'].replace("\xa0", " "),
                                data['Content7'].replace("\xa0", " "), data['Content8'].replace("\xa0", " "),
                                data['Content9'].replace("\xa0", " ")]
                    Content = filter(no, Contents)
                    s1 = ""
                    for j in Content:
                        s1 = s1 + j + "\n"
                    s = "("+str(k+1)+")、"+Question + "\n" + s1
                    question_lst.append(s)
                    c += 1
                    k += 1
                i += 1
            else:
                break

        f = open("English_choice_2.json", encoding='utf-8')  # 返回一个文件对象
        line = json.load(f)  # 调用文件的 readline()方法
        length = len(line)
        i = 0
        c = 0
        while i < length:
            if c < len(n2):
                if i == n2[c]:
                    data = line[i]
                    Question = data['Question']
                    Choices = data['Choices']
                    s2 = "("+str(k+1)+")、"+Question + "\n" + Choices
                    question_lst.append(s2)
                    c += 1
                    k += 1
                i += 1
            else:
                break

        return question_lst

    def Sentence_Question(self):
        number = self.number3
        f = open("English_sentence.json", encoding='utf-8')  # 返回一个文件对象
        line = json.load(f)  # 调用文件的 readline()方法
        length = len(line)
        a = get_random_list(0, 55, number)
        n = sorted(a)
        question_lst = ['连词成句部分']
        i = 0
        c = 0
        k=0
        while i < length:
            if c < len(n):
                if i == n[c]:
                    data = line[i]
                    Question = data['Question']
                    s = "("+str(k+1)+")、"+Question + "\n"
                    question_lst.append(s)
                    c += 1
                    k += 1
                i += 1
            else:
                break

        return question_lst

    def Writing(self):
        number = self.number4
        f = open("English_write.json", encoding='utf-8')  # 返回一个文件对象
        line = json.load(f)  # 调用文件的 readline()方法
        length = len(line)
        a = get_random_list(0, 18, number)
        n = sorted(a)
        question_lst = ['写作部分']
        i = 0
        c = 0
        while i < length:
            if c < len(n):
                if i == n[c]:
                    data = line[i]
                    Question = data['Question']
                    Contents = [data['Content1'].replace("\xa0", ""), data['Content2'].replace("\xa0", " "),
                                data['Content3'].replace("\xa0", " "), data['Content4'].replace("\xa0", " "),
                                data['Content5'].replace("\xa0", " "), data['Content6'].replace("\xa0", " ")]
                    Content = filter(no, Contents)
                    str = ""
                    for j in Content:
                        str = str + j + "\n"
                    s = Question + "\n" + str
                    question_lst.append(s)
                    c += 1
                i += 1
            else:
                break

        return question_lst


def main():
    number=30
    number2 = 10
    number3 = 15
    number4 = 1
    spelling=0
    intertranslation=0
    is_English=0
    is_Chnese=0
    is_Mixed=0
    c=[]
    g=EnglishGenerate(spelling,intertranslation,is_English,is_Chnese, is_Mixed, number,number2,number3,number4)
    l=g.Spelling()
    c.append(l)
    l2=g.Choice_Question()
    l3=g.Sentence_Question()
    l4=g.Writing()
    t = ['小学生英语试卷3']
    pp = PrintPreview(l=c, l2=l2, l3=l3, l4=l4, m1=0, m2=0, m3=0, m4=1, tit=t, col=3,
                      subtitle="姓名：__________ 日期：____月____日 时间：________ 对题：____道", )
    pp.produce()

if __name__ == '__main__':
    main()









        
