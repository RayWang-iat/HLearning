'''
这是处理数学题目的类
'''

import random
import os
import xlrd

from docx import Document  # 引入docx类生成docx文档
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Math_xz:
    '''
    这是处理选择题的类
    '''
    data = []
    list_q = []

    def load(self):
        workbook = xlrd.open_workbook('数学选择题.xlsx')
        booksheet = workbook.sheet_by_name('Sheet1')
        for i in range(booksheet.nrows):
            xz = []
            for j in range(booksheet.ncols):
                cell = booksheet.cell(i,j)
                xz.append(cell.value)
            self.data.append(xz)

    def get_q(self,Q_num,T_num):
        for i in range(T_num):
            list_every_t = []
            list_rand_q = random.sample(range(0,len(self.data)),Q_num)
            for j in list_rand_q:
                list_every_t.append(self.data[j])
            self.list_q.append(list_every_t)

    def deal_q(self):
        list_temp = []
        for i in self.list_q:
            list_temp_one = []
            for j in i:
                list_temp_every = []
                list_temp_every.append(j[0])
                temp = str(j[1]) + "      " + str(j[2]) + "      " + str(j[3]) + "      " + str(j[4])
                list_temp_every.append(temp)
                list_temp_one.append((list_temp_every))
            list_temp.append(list_temp_one)
        self.list_q = list_temp

    def print_q(self):
        self.deal_q()
        for i in range(len(self.list_q)):
            for j in self.list_q[i]:
                print(j)
            print("\n\n\n\n\n")


    def print(self):
        for i in self.data:
            print(i)

class Math_pd:

    list_data = []
    list_q = []

    def load_pd(self):
        workbook = xlrd.open_workbook('数学判断题.xlsx')
        booksheet = workbook.sheet_by_name('Sheet1')
        for i in range(booksheet.nrows):
            list_temp = []
            for j in range(booksheet.ncols):
                cell = booksheet.cell(i,j)
                list_temp.append(cell.value)
            self.list_data.append(list_temp)

    def get_q(self,Q_num,T_num):
        for i in range(T_num):
            list_temp = []
            list_rand = random.sample(range(0,len(self.list_data)),Q_num)
            for j in list_rand:
                list_temp.append(self.list_data[j])
            self.list_q.append(list_temp)

    def print_list_data(self):
        for i in self.list_data:
            print(i)

    def print_list_q(self):
        for i in self.list_q:
            for j in i:
                print(j)
            print("\n\n")

class Math_ss:

    list_data = []
    list_q = []

    def load_ss(self):
        workbook = xlrd.open_workbook('数学计算题.xlsx')
        booksheet = workbook.sheet_by_name('Sheet1')
        for i in range(booksheet.nrows):
            list_temp = []
            for j in range(booksheet.ncols):
                cell = booksheet.cell(i,j)
                list_temp.append(cell.value)
            self.list_data.append(list_temp)

    def get_q(self,Q_num,T_num):
        for i in range(T_num):
            list_temp = []
            list_rand = random.sample(range(0,len(self.list_data)),Q_num)
            for j in list_rand:
                list_temp.append(self.list_data[j])
            self.list_q.append(list_temp)

    def print_list_q(self):
        for i in self.list_q:
            for j in i:
                print(j)
            print("\n\n\n")


    def print_list_data(self):
        for i in self.list_data:
            print(i)

class Math_tk:

    list_data = []
    list_q = []

    def load_tk(self):
        workbook = xlrd.open_workbook('数学填空题.xlsx')
        booksheet = workbook.sheet_by_name('Sheet1')
        for i in range(booksheet.nrows):
            list_temp = []
            for j in range(booksheet.ncols):
                cell = booksheet.cell(i,j)
                list_temp.append(cell.value)
            self.list_data.append(list_temp)

    def get_q(self,Q_num,T_num):
        for i in range(T_num):
            list_temp = []
            list_rand = random.sample(range(0,len(self.list_data)),Q_num)
            for j in list_rand:
                list_temp.append(self.list_data[j])
            self.list_q.append(list_temp)

    def print_list_q(self):
        for i in self.list_q:
            for j in i:
                print(j)
            print("\n\n\n")


    def print_list_data(self):
        for i in self.list_data:
            print(i)

class Math_yy:

    list_data = []
    list_q = []

    def load(self):
        workbook = xlrd.open_workbook('数学应用题.xlsx')
        bookesheet = workbook.sheet_by_name('Sheet1')
        for i in range(bookesheet.nrows):
            list_temp = []
            for j in range(bookesheet.ncols):
                cell = bookesheet.cell(i,j)
                list_temp.append(cell.value)
            self.list_data.append(list_temp)

    def get_q(self,Q_num,T_num):
        for i in range(T_num):
            list_temp = []
            list_rand = random.sample(range(0,len(self.list_data)),Q_num)
            for j in list_rand:
                list_temp.append(self.list_data[j])
            self.list_q.append(list_temp)

    def print_list_q(self):
        for i in self.list_data:
            for j in i:
                print(j)
            print("\n")

    def print_list_data(self):
        for i in self.list_data:
            print(i)


class PrintPreview:

    def __init__(self,l_xz,l_tk, l_pd,l_ss,l_yy,title_main, subtitle, col=1, tsize=26, subsize=11, csize=16,
                 # 默认输出文件地址为项目根目录
                 docxpath=os.path.join(os.path.dirname(os.path.abspath(__file__)), "")):
        '''
        :param l: list 需要打印的口算题列表
        :param tit: list 口算页标题
        :param subtitle str 小标题
        :param col: int 列数
        :param tsize: int 标题字号
        :param csize: int 口算题字号
        :param docxpath str 保存路径

        '''
        self.xz = l_xz
        self.tk = l_tk
        self.pd = l_pd
        self.ss = l_ss
        self.yy = l_yy
        self.p_title_main = title_main
        self.p_subtitle = subtitle
        self.p_column = col
        self.p_title_size = tsize
        self.p_subtitle_size = subsize
        self.p_content_siae = csize
        self.docxpath = docxpath
        print(self.docxpath)

    def create_psmdocx(self,l_xz,l_tk, l_pd,l_ss,l_yy, docxname):

        '''
        :param l list 一组题库
        :param title str 页面标题
        :param docxname  str 题库保存文件名
        :return: none
        '''
        # if (title == ''):
        #     page_title = '小学生口算题'
        # else:
        #     page_title = title
        page_title = self.p_title_main
        p_docx = Document()  # 创建一个docx文档
        p_docx.styles['Normal'].font.name = u'Times'  # 可换成word里面任意字体
        p = p_docx.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        run = p.add_run(page_title)
        run.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(self.p_title_size)  # 字体大小设置，和word里面的字号相对应

        sp = p_docx.add_paragraph()
        sp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        srun = sp.add_run(self.p_subtitle)
        srun.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun.font.size = Pt(self.p_subtitle_size)  # 字体大小设置，和word里面的字号相对应

        # 判断需要用到的行数
        # if (len(l) % self.p_column):
        #     rs = len(l) // self.p_column + 2
        # else:
        #     rs = len(l) // self.p_column + 1
        len_xz = len(l_xz)
        len_tk = len(l_tk)
        len_ss = len(l_ss)
        len_pd = len(l_pd)
        len_yy = len(l_yy)
        # print(rs)

        # 将口算题添加到docx表格中
        jishu = 1  # 计数器
        table = p_docx.add_table(rows=len_xz * 2 + 1 + len_tk + 1 + len_pd + 1 + len_ss + 1 + 3*len_yy + 1, cols=1)
        if(len_xz > 0):
            table.rows[0].cells[0].text = str(jishu) + "、单项选择："
            jishu += 1
            choose_num = 1
            for i in range(len_xz):
                list_temp = l_xz[i]
                str_choose = "(" + str(choose_num) + ")"
                choose_num += 1
                str_choose += list_temp[0]
                table.rows[2*i+1].cells[0].text = str_choose
                str_xx = ""
                str_xx += list_temp[1]
                table.rows[2*i+2].cells[0].text = str_xx

        if(len_pd > 0):

            table.rows[2*len_xz+1].cells[0].text = str(jishu) + "、判断题："
            jishu += 1
            # list_poem_temp = []
            q_num = 1
            for i in range(len(l_pd)):
                list_temp = l_pd[i]
                string = ""
                string = "(" + str(q_num) + ")"
                string += list_temp[0]
                table.rows[2*len_xz+2+i].cells[0].text = string
                q_num += 1


        if(len_tk > 0):
            table.rows[2*len_xz + 1 + len_pd + 1].cells[0].text = str(jishu) + "、填空题："
            jishu += 1
            s_q = 1
            list_end = []
            for i in range(len(l_tk)):
                list_temp = l_tk[i]
                str_sen = "(" + str(s_q) + ")"
                str_sen += list_temp[0]
                table.rows[2 * len_xz + 1 + len_pd + 1 + 1 + i].cells[0].text = str_sen
                s_q += 1


        if(len_ss > 0):
            table.rows[2 * len_xz + 1 + len_pd + 1 + len_tk + 1].cells[0].text = str(jishu) + "、计算题"
            jishu += 1
            q_num = 1
            for i in range(len(l_ss)):
                list_temp = l_ss[i]
                str_ss = "(" + str(q_num) + ")"
                str_ss += list_temp[0]
                table.rows[2 * len_xz +1 + len_pd + 1 + len_tk + 1 + 1 + i].cells[0].text = str_ss
                q_num += 1

        if(len_yy > 0):
            table.rows[2 * len_xz + 1 + len_pd + 1 + len_tk + 1 + len_ss + 1].cells[0].text = str(jishu) + "、应用题"
            jishu += 1
            q_num = 1
            for i in range(len(l_yy)):
                list_temp = l_yy[i]
                str_yy = "(" + str(q_num) + ")"
                str_yy += list_temp[0]
                table.rows[2 * len_xz +1 + len_pd + 1 + len_tk + 1 + len_ss + 1 + 1 + 3*i].cells[0].text = str_yy
                q_num += 1

        table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        table.style.font.size = Pt(self.p_content_siae)  # 字体大小设置，和word里面的字号相对应

        print(self.docxpath + '{}.docx'.format(docxname))
        p_docx.save(self.docxpath + '{}.docx'.format(docxname))  # 输出docx

    def produce(self):
        k = 1
        # for l,l_s in self.p_list,self.sen_list:
        #     self.create_psmdocx(l,l_s, "小学语文古诗词填空"+ str(k))
        #     k = k + 1
        for i in range(len(self.xz)):
            self.create_psmdocx(self.xz[i],self.tk[i],self.pd[i],self.ss[i],self.yy[i],"小学数学练习题" + str(k))
            k += 1


if __name__ == '__main__':
    xz = Math_xz()
    xz.load()
    xz.get_q(5,4)
    xz.deal_q()
    # xz.print_q()
    # xz.print()
    pd = Math_pd()
    pd.load_pd()
    pd.get_q(5,4)
    # pd.print_list_data()
    # pd.print_list_q()
    ss = Math_ss()
    ss.load_ss()
    ss.get_q(10,4)
    # ss.print_list_q()

    tk = Math_tk()
    tk.load_tk()
    tk.get_q(5,4)
    # tk.print_list_q()

    yy = Math_yy()
    yy.load()
    # yy.print_list_data()
    yy.get_q(5,4)
    # yy.print_list_q()
    pp = PrintPreview(xz.list_q,tk.list_q,pd.list_q,ss.list_q,yy.list_q,"小学数学练习题","姓名：__________ 日期：____月____日 时间：________ 对题：____道")
    pp.produce()



