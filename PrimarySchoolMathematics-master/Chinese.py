'''
生成语文试题的类
'''
import random
import os
import xlrd

from docx import Document  # 引入docx类生成docx文档
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Chinese:
    '''
    生成的试卷数：T_num
    每张试卷的题数：Q_num
    生成试卷的模式：Model：0-随机、1-写出诗词的名字、2-写出诗词的作者，3-写出诗词的前一句，4-写出诗词的后一句
    存放从文件读取的古诗词：list_poem； type：list
    存放最后生成的试题：list_q
    '''
    T_num = None
    Q_num = None
    Model = None
    list_poem = []
    list_q = []




    def Load_poem(self):
        '''
        读取poem文件，存放到list_poem中
        :return:
        '''
        f = open("poem.txt",encoding='utf-8')
        list_temp = []
        str = ""
        while 1:
            line = f.readline()
            if not line:
                break
            pass
            if line.startswith("<"):
                line_temp = "《"
                line_temp += line[1:len(line)-2]
                line_temp += "》"
                list_temp.append(line_temp)
                str = ""
            elif line.startswith("&"):
                line_temp = line[1:len(line)-1]
                list_temp.append(line_temp)
            else:
                str = str + line
                str = str[0:len(str)-1]
                poem_main_temp = str.split('|')
                poem_main = []
                for i in poem_main_temp:
                    sentence = i.split('#')
                    poem_main.append(sentence)
                list_temp.append(poem_main)
                self.list_poem.append(list_temp)
                list_temp = []
        f.close()

    def Model1(self,T_num,Q_num):
        '''
        生成一组写出古诗词题目的试题
        :param T_num: 试卷数
        :param Q_num: 每张试卷的题数
        :return: None
        '''
        self.T_num = T_num
        self.Q_num = Q_num
        for i in range(T_num):
            list_q_temp = []
            L1 = random.sample(range(1,len(self.list_poem)-1),self.Q_num)
            for j in L1:
                temp_copy = self.list_poem.copy()
                temp1 = temp_copy[j]
                temp = []
                temp.append("_____________")
                num = len(temp1)
                for k in range(num-1):
                    temp.append(temp1[k+1])
                list_q_temp.append(temp)
            self.list_q.append(list_q_temp)

    def Model2(self,T_num,Q_num):
        '''
        生成一组写出作者是谁的题目
        :param T_num: 试卷的数量
        :param Q_num: 每张试卷里面的题目量
        :return: None
        '''
        self.T_num = T_num
        self.Q_num = Q_num
        for i in range(T_num):
            list_q_temp = []
            L1 = random.sample(range(1,len(self.list_poem)-1),self.Q_num)
            for j in L1:
                temp_copy = self.list_poem.copy()
                temp1 = temp_copy[j]
                temp = []
                temp.append(temp1[0])
                temp.append("_____________")
                num = len(temp1)
                for k in range(num-2):
                    temp.append(temp1[k+2])
                list_q_temp.append(temp)
            self.list_q.append(list_q_temp)

    def Model3(self,T_num,Q_num):
        '''
        生成一组根据后面的句子写出前面的句子的题目
        :param T_num: 试卷的数量
        :param Q_num: 每张试卷里面的题目量
        :return: None
        '''
        for i in range(T_num):
            list_q_temp = []
            L1 = random.sample(range(0,len(self.list_poem)-1),Q_num)
            for i in L1:
                temp_copy = self.list_poem.copy()
                temp_onepoem = temp_copy[i]
                temp_blank = []
                temp_blank.append(temp_onepoem[0])
                temp_blank.append(temp_onepoem[1])
                temp_poem_main = temp_onepoem[2]
                temp_poem_main_sentence = temp_poem_main[random.randint(1,len(temp_poem_main))-1]
                temp_blank.append("_______________")
                num = len(temp_poem_main_sentence)
                for k in range(num-1):
                    temp_blank.append(temp_poem_main_sentence[k+1])
                list_q_temp.append(temp_blank)
            self.list_q.append(list_q_temp)

    def Model4(self,T_num,Q_num):
        '''
        生成一组根据前面的句子写出后面的句子的题目
        :param T_num: 试卷的数量
        :param Q_num: 每张试卷里面的题目量
        :return: None
        '''
        for i in range(T_num):
            list_q_temp = []
            L1 = random.sample(range(1,len(self.list_poem)-1),Q_num)
            for i in L1:
                temp_copy = self.list_poem.copy()
                temp_onepoem = temp_copy[i]
                temp_blank = []
                temp_blank.append(temp_onepoem[0])
                temp_blank.append(temp_onepoem[1])
                temp_poem_main = temp_onepoem[2]
                temp_poem_main_sentence = temp_poem_main[random.randint(1,len(temp_poem_main))-1]
                temp_blank.append(temp_poem_main_sentence[0])
                temp_blank.append("_______________")
                num = len(temp_poem_main_sentence)
                for k in range(num - 2):
                    temp_blank.append(temp_poem_main_sentence[k + 2])
                list_q_temp.append(temp_blank)
            self.list_q.append(list_q_temp)

    def model1_one_q(self,Q_num):
        '''
        这是生成随机题目类型下面的子函数，这个函数是根据需要题目的数量生成题目（写诗名）
        :param Q_num:需要生成题目量
        :return:
        '''
        list_q_temp = []
        L1 = random.sample(range(0, len(self.list_poem)-1), Q_num)
        for j in L1:
            temp_copy = self.list_poem.copy()
            temp1 = temp_copy[j]
            temp = []
            temp.append("_____________")
            num = len(temp1)
            for k in range(num-1):
                temp.append(temp1[k+1])
            list_q_temp.append(temp)
        return list_q_temp

    def model2_one_q(self,Q_num):
        '''
        这是生成随机题目类型下面的子函数，这个函数是根据需要题目的数量生成题目（写作者）
        :param Q_num:需要生成题目量
        :return:
        '''
        list_q_temp = []
        print(len(self.list_poem),Q_num)
        L1 = random.sample(range(0, len(self.list_poem)-1), Q_num)
        for j in L1:
            temp_copy = self.list_poem.copy()
            temp1 = temp_copy[j]
            temp = []
            temp.append(temp1[0])
            temp.append("_____________")
            num = len(temp1)
            for k in range(num-2):
                temp.append(temp1[k+2])
            list_q_temp.append(temp)
        return list_q_temp

    def model3_one_q(self,Q_num):
        '''
        这是生成随机题目类型下面的子函数，这个函数是根据需要题目的数量生成题目（根据后句写前句）
        :param Q_num:需要生成题目量
        :return:
        '''
        list_q_temp = []
        L1 = random.sample(range(0, len(self.list_poem)-1), Q_num)
        for i in L1:
            temp_copy = self.list_poem.copy()
            temp_onepoem = temp_copy[i]
            temp_blank = []
            temp_blank.append(temp_onepoem[0])
            temp_blank.append(temp_onepoem[1])
            temp_poem_main = temp_onepoem[2]
            temp_poem_main_sentence = temp_poem_main[random.randint(0, len(temp_poem_main)-1)]
            temp_blank.append("_______________")
            num = len(temp_poem_main_sentence)
            for k in range(num-1):
                temp_blank.append(temp_poem_main_sentence[k+1])
            list_q_temp.append(temp_blank)
        return list_q_temp

    def model4_one_q(self,Q_num):
        '''
        这是生成随机题目类型下面的子函数，这个函数是根据需要题目的数量生成题目（根据前句写后句）
        :param Q_num:需要生成题目量
        :return:
        '''
        list_q_temp = []
        L1 = random.sample(range(0, len(self.list_poem)-1), Q_num)
        for i in L1:
            temp_copy = self.list_poem.copy()
            temp_onepoem = temp_copy[i]
            temp_blank = []
            temp_blank.append(temp_onepoem[0])
            temp_blank.append(temp_onepoem[1])
            temp_poem_main = temp_onepoem[2]
            temp_poem_main_sentence = temp_poem_main[random.randint(1, len(temp_poem_main)) - 1]
            temp_blank.append(temp_poem_main_sentence[0])
            temp_blank.append("_______________")
            num = len(temp_poem_main_sentence)
            for k in range(num - 2):
                temp_blank.append(temp_poem_main_sentence[k + 2])
            list_q_temp.append(temp_blank)
        return list_q_temp

    def Model0(self,T_num,Q_num):
        '''
        这是随机生成各种各样类型的题目的函数，即每张卷子里面各种类型的题目都有
        :param T_num: 试卷的数量
        :param Q_num: 每张试卷的题目的数量
        :return:
        '''
        for i in range(T_num):
            list_q_temp = []
            one = 0
            two = 0
            three = 0
            four = 0
            for j in range(Q_num):
                num = random.randint(1,4)
                if(num == 0):
                    one +=1
                elif(num == 1):
                    two += 1
                elif(num == 2):
                    three += 1
                else:
                    four += 1
            poem_temp = self.model1_one_q(one)
            for one_temp in poem_temp:
                list_q_temp.append(one_temp)
            poem_temp = self.model2_one_q(two)
            for two_temp in poem_temp:
                list_q_temp.append(two_temp)
            poem_temp = self.model3_one_q(three)
            for three_temp in poem_temp:
                list_q_temp.append(three_temp)
            poem_temp = self.model4_one_q(four)
            for four_temp in poem_temp:
                list_q_temp.append(four_temp)
            self.list_q.append(list_q_temp)

    def print_poem(self):
        '''
        打印加载到古诗词列表中的诗句，仅用来调试与预览
        :return:
        '''
        self.Load_poem()
        for i in self.list_poem:
            print(i)

    def print_Model(self):
        '''
        打印试卷里面的试题，用来测试与预览
        :return:
        '''
        for i in self.list_q:
            print(i)

class Chinese_article:
    list_article = []
    list_question = []
    def load_article(self):
        workbook = xlrd.open_workbook('Chinese_article.xlsx')
        booksheet = workbook.sheet_by_name('Sheet1')
        for row in range(booksheet.nrows):
            sentence = []
            for col in range(booksheet.ncols):
                cel = booksheet.cell(row, col)
                val = cel.value
                sentence.append(val)
            self.list_article.append(sentence)

    def get_question(self,num_T):
        for i in range(num_T):
            num = random.randint(0,len(self.list_article)-1)
            self.list_question.append(self.list_article[num])

    def print_article(self):
        for i in range(len(self.list_question)):
            print(self.list_question)

    def print_all(self):
        for i in range(len(self.list_article)):
            print(self.list_article[i])


class Chinese_Choose:
    list_choose = []
    list_question = []

    def load_choose(self):
        workbook = xlrd.open_workbook('小学语文.xlsx')
        booksheet = workbook.sheet_by_name('Sheet1')
        for row in range(booksheet.nrows):
            cel = booksheet.cell(row, 0)
            val = cel.value
            self.list_choose.append(val)

    def get_question_choose(self, num_T, num_Q):
        for i in range(num_T):
            list_sen_temp = []
            random_list = random.sample(range(0, len(self.list_choose)), num_Q)
            for j in random_list:
                list_sen_temp.append(self.list_choose[j])
            self.list_question.append(list_sen_temp)

    def print_all(self):
        for i in range(len(self.list_choose)):
            print(self.list_choose[i])

class Chinese_sentence:

    list_sentence = []
    list_ques = []

    def load(self):
        workbook = xlrd.open_workbook('Chinese_sentence.xlsx')
        booksheet = workbook.sheet_by_name('Sheet1')
        for row in range(booksheet.nrows):
            sentence = []
            for col in range(booksheet.ncols):
                cel = booksheet.cell(row,col)
                val = cel.value
                sentence.append(val)
            self.list_sentence.append(sentence)
        # print(self.list_sentence)

    def getQuestion(self,num_T,num_Q):
        for i in range(num_T):
            list_sen_temp = []
            random_list = random.sample(range(0,len(self.list_sentence)),num_Q)
            for j in random_list:
                list_sen_temp.append(self.list_sentence[j])
            self.list_ques.append(list_sen_temp)

    # def print_Question(self):
    #     for i in range(len(self.list_ques)):
    #         list_temp = self.list_sentence[i]
    #         print(list_temp,end="\n")


class PrintPreview:

    def __init__(self,l_c,l_a, l,l_s,title_main, subtitle, col=1, tsize=26, subsize=11, csize=16,
                 # 默认输出文件地址为项目根目录
                 docxpath=os.path.join(os.path.dirname(os.path.abspath(__file__)), "")):
        '''
        :param l: list 需要打印的口算题列表
        :param tit: list 口算页标题
        :param subtitle str 小标题
        :param col: int 列数
        :param tsize: int 标题字号
        :param csize: int 口算题字号
        :param docxpath str 保存路径

        '''
        self.p_choose = l_c
        self.p_article = l_a
        self.p_title_main = title_main
        self.p_list = l
        self.sen_list = l_s
        self.p_subtitle = subtitle
        self.p_column = col
        self.p_title_size = tsize
        self.p_subtitle_size = subsize
        self.p_content_siae = csize
        self.docxpath = docxpath
        print(self.docxpath)

    def create_psmdocx(self,l_c,l_a, l,l_s, docxname):
        '''
        :param l list 一组题库
        :param title str 页面标题
        :param docxname  str 题库保存文件名
        :return: none
        '''
        # if (title == ''):
        #     page_title = '小学生口算题'
        # else:
        #     page_title = title
        page_title = self.p_title_main
        p_docx = Document()  # 创建一个docx文档
        p_docx.styles['Normal'].font.name = u'Times'  # 可换成word里面任意字体
        p = p_docx.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        run = p.add_run(page_title)
        run.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(self.p_title_size)  # 字体大小设置，和word里面的字号相对应

        sp = p_docx.add_paragraph()
        sp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        srun = sp.add_run(self.p_subtitle)
        srun.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun.font.size = Pt(self.p_subtitle_size)  # 字体大小设置，和word里面的字号相对应

        # 判断需要用到的行数
        # if (len(l) % self.p_column):
        #     rs = len(l) // self.p_column + 2
        # else:
        #     rs = len(l) // self.p_column + 1
        rs = len(l)
        len_c = len(l_c)
        len_s = len(l_s)
        len_a = len(l_a)
        # print(rs)

        # 将口算题添加到docx表格中
        jishu = 1  # 计数器
        table = p_docx.add_table(rows=len_c + 1 + rs + 1 + 2 * len_s + 1 + 3, cols=1)
        if(len_c > 0):
            table.rows[0].cells[0].text = str(jishu) + "、单项选择："
            jishu += 1
            choose_num = 1
            for i in range(len_c):
                str_choose = "(" + str(choose_num) + ")"
                choose_num += 1
                str_choose += l_c[i]
                table.rows[i+1].cells[0].text = str_choose
        if(rs > 0):

            table.rows[len_c+1].cells[0].text = str(jishu) + "、填写下面诗句中所缺的内容（诗句，作者，诗名）："
            jishu += 1
            list_poem_temp = []
            q_num = 1
            for i in range(len(l)):
                string = ""
                string = "(" + str(q_num) + ")"
                poem_temp = l[i]
                for m in range(len(poem_temp)-2):
                    poem_main = poem_temp[m+2]
                    # print("poem_main",poem_main)
                    for k in range(len(poem_main)):
                        poem_sentence = poem_main[k]
                        for j in range(len(poem_sentence)):
                            string += poem_sentence[j]
                string += "    "
                string += poem_temp[0]
                string += "    "
                string += poem_temp[1]
                list_poem_temp.append(string)
                q_num += 1
            for i in range(rs):
                row_cells = table.rows[len_c+i+2].cells
                row_cells[0].text = list_poem_temp[i]
        if(len_s > 0):
            table.rows[len_c+1+rs+1].cells[0].text = str(jishu) + "、根据要求完成下列句子："
            jishu += 1
            s_q = 1
            list_end = []
            for i in range(len(l_s)):
                str_sen = "(" + str(s_q) + ")"
                list_sen_temp = l_s[i]
                str_sen += list_sen_temp[1] + "(" + list_sen_temp[0] + ")"
                str_sen2 = list_sen_temp[2]
                list_end.append(str_sen)
                list_end.append(str_sen2)
                s_q += 1
            for i in range(len(list_end)):
                row_cells = table.rows[len_c + 1 + rs + 1 + i + 1].cells
                row_cells[0].text = list_end[i]
        if(len_a > 0):
            table.rows[len_c + 1 + rs + 1 + 2 * len_s + 1].cells[0].text = str(jishu) + "、" + l_a[0]
            table.rows[len_c + 1 + rs + 1 + 2 * len_s + 2].cells[0].text = l_a[1]
            table.rows[len_c + 1 + rs + 1 + 2 * len_s + 3].cells[0].text = l_a[2]
        table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        table.style.font.size = Pt(self.p_content_siae)  # 字体大小设置，和word里面的字号相对应

        print(self.docxpath + '{}.docx'.format(docxname))
        p_docx.save(self.docxpath + '{}.docx'.format(docxname))  # 输出docx

    def produce(self):
        k = 1
        # for l,l_s in self.p_list,self.sen_list:
        #     self.create_psmdocx(l,l_s, "小学语文古诗词填空"+ str(k))
        #     k = k + 1
        for i in range(len(self.p_list)):
            self.create_psmdocx(self.p_choose[i],self.p_article[i],self.p_list[i],self.sen_list[i],"小学语文练习题" + str(k))
            k += 1


if __name__ == "__main__":
    ch = Chinese()
    ch.Load_poem()
    # ch.print_poem()
    ch.Model0(5,10)      #测试S
    # ch.print_Model()
    ch_choose = Chinese_Choose()
    ch_choose.load_choose()
    ch_choose.get_question_choose(5,5)
    ch_sentence = Chinese_sentence()
    ch_sentence.load()
    ch_sentence.getQuestion(5,5)
    ch_article = Chinese_article()
    ch_article.load_article()
    ch_article.get_question(5)
    pp = PrintPreview(ch_choose.list_question,ch_article.list_question,ch.list_q,ch_sentence.list_ques,"小学练习题","姓名：__________ 日期：____月____日 时间：________ 对题：____道")
    pp.produce()





